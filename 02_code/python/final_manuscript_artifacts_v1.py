from __future__ import annotations

import json
import math
import shutil
import textwrap
from html import escape
from dataclasses import dataclass
from pathlib import Path

import numpy as np
import pandas as pd
from PIL import Image, ImageDraw, ImageFont


RUN_ID = "20260621_final_manuscript_artifacts_v1"
PROJECT_ROOT = Path(__file__).resolve().parents[2]
OUTPUT_ROOT = PROJECT_ROOT / "03_outputs" / "final_artifacts" / RUN_ID
TABLE_DIR = OUTPUT_ROOT / "tables"
FIGURE_DIR = OUTPUT_ROOT / "figures"
QA_DIR = OUTPUT_ROOT / "qa"
STAGING_DIR = OUTPUT_ROOT / "staging"


COLORS = {
    "navy": "#183A59",
    "blue": "#2D6C8E",
    "blue_light": "#EAF3F7",
    "green": "#2F6B57",
    "green_light": "#EAF4EF",
    "purple": "#6A5A8E",
    "purple_light": "#F0EDF7",
    "gold": "#B8892F",
    "gold_light": "#FBF2DB",
    "gray": "#68737D",
    "gray_light": "#F4F6F7",
    "red": "#9A3F3F",
}


def ensure_dirs() -> None:
    for directory in [OUTPUT_ROOT, TABLE_DIR, FIGURE_DIR, QA_DIR, STAGING_DIR]:
        directory.mkdir(parents=True, exist_ok=True)


def read_csv_auto(path: Path) -> pd.DataFrame:
    last_error: Exception | None = None
    for encoding in ["utf-8-sig", "utf-16", "utf-16-le", "gbk", "latin1"]:
        try:
            return pd.read_csv(path, encoding=encoding)
        except Exception as exc:  # pragma: no cover - diagnostic fallback
            last_error = exc
    raise RuntimeError(f"Could not read {path}") from last_error


def source(path: str) -> Path:
    return PROJECT_ROOT / path


def fmt_int(x: float | int | str | None) -> str:
    if x is None or (isinstance(x, float) and math.isnan(x)):
        return "-"
    return f"{int(round(float(x))):,}"


def fmt_float(x: float | int | None, digits: int = 3) -> str:
    if x is None or (isinstance(x, float) and math.isnan(x)):
        return "-"
    return f"{float(x):.{digits}f}"


def fmt_p(p: float | None) -> str:
    if p is None or (isinstance(p, float) and math.isnan(p)):
        return "-"
    if p < 0.001:
        return "<0.001"
    return f"{p:.3f}"


def normal_two_sided_p(z: float) -> float:
    return math.erfc(abs(float(z)) / math.sqrt(2.0))


def mann_whitney_p(x: pd.Series, group: pd.Series) -> float | None:
    values = pd.to_numeric(x, errors="coerce")
    groups = pd.to_numeric(group, errors="coerce")
    valid = values.notna() & groups.notna()
    values = values[valid]
    groups = groups[valid]
    x0 = values[groups.eq(0)].to_numpy()
    x1 = values[groups.eq(1)].to_numpy()
    n0, n1 = len(x0), len(x1)
    if n0 == 0 or n1 == 0:
        return None
    combined = pd.Series(np.concatenate([x0, x1]))
    ranks = combined.rank(method="average").to_numpy()
    r1 = ranks[n0:].sum()
    u1 = r1 - n1 * (n1 + 1) / 2.0
    mean_u = n0 * n1 / 2.0
    n = n0 + n1
    tie_counts = combined.value_counts().to_numpy()
    tie_term = np.sum(tie_counts ** 3 - tie_counts)
    var_u = n0 * n1 / 12.0 * ((n + 1) - tie_term / (n * (n - 1))) if n > 1 else 0
    if var_u <= 0:
        return None
    correction = 0.5 if u1 > mean_u else -0.5
    z = (u1 - mean_u - correction) / math.sqrt(var_u)
    return normal_two_sided_p(z)


def comb(n: int, k: int) -> int:
    if k < 0 or k > n:
        return 0
    return math.comb(int(n), int(k))


def fisher_exact_two_sided(a: int, b: int, c: int, d: int) -> float:
    row1, row2 = a + b, c + d
    col1, total = a + c, a + b + c + d

    def prob(x: int) -> float:
        return comb(row1, x) * comb(row2, col1 - x) / comb(total, col1)

    observed = prob(a)
    lo = max(0, col1 - row2)
    hi = min(row1, col1)
    p = sum(prob(x) for x in range(lo, hi + 1) if prob(x) <= observed + 1e-12)
    return min(1.0, p)


def chi_square_2x2_p(a: int, b: int, c: int, d: int) -> float | None:
    total = a + b + c + d
    if total == 0:
        return None
    row1, row2 = a + b, c + d
    col1, col2 = a + c, b + d
    expected = [
        row1 * col1 / total,
        row1 * col2 / total,
        row2 * col1 / total,
        row2 * col2 / total,
    ]
    if any(e <= 0 for e in expected):
        return None
    observed = [a, b, c, d]
    chi2 = sum((o - e) ** 2 / e for o, e in zip(observed, expected))
    return math.erfc(math.sqrt(chi2 / 2.0))


def binary_p(x: pd.Series, group: pd.Series) -> float | None:
    values = pd.to_numeric(x, errors="coerce")
    groups = pd.to_numeric(group, errors="coerce")
    valid = values.notna() & groups.notna()
    values = values[valid]
    groups = groups[valid]
    if values.empty or groups.nunique() < 2:
        return None
    a = int(((groups == 0) & (values == 1)).sum())
    b = int(((groups == 0) & (values == 0)).sum())
    c = int(((groups == 1) & (values == 1)).sum())
    d = int(((groups == 1) & (values == 0)).sum())
    total = a + b + c + d
    if total == 0:
        return None
    row1, row2 = a + b, c + d
    col1, col2 = a + c, b + d
    expected = [
        row1 * col1 / total,
        row1 * col2 / total,
        row2 * col1 / total,
        row2 * col2 / total,
    ]
    if any(e < 5 for e in expected):
        return fisher_exact_two_sided(a, b, c, d)
    return chi_square_2x2_p(a, b, c, d)


def fmt_or_ci(row: pd.Series) -> str:
    return f"{row['or']:.3f} ({row['ci_low']:.3f}-{row['ci_high']:.3f})"


def med_iqr(series: pd.Series, digits: int = 1) -> str:
    vals = pd.to_numeric(series, errors="coerce").dropna()
    if vals.empty:
        return "-"
    q25, med, q75 = vals.quantile([0.25, 0.5, 0.75]).to_numpy()
    return f"{med:.{digits}f} [{q25:.{digits}f}-{q75:.{digits}f}]"


def n_pct(series: pd.Series, total: int | None = None, digits: int = 1) -> str:
    vals = pd.to_numeric(series, errors="coerce")
    denom = int(vals.notna().sum()) if total is None else int(total)
    if denom == 0:
        return "-"
    n = int((vals == 1).sum())
    return f"{n:,} ({n / denom * 100:.{digits}f})"


def n_pct_from_counts(n: int, denom: int, digits: int = 1) -> str:
    if denom == 0:
        return "-"
    return f"{int(n):,} ({int(n) / int(denom) * 100:.{digits}f})"


def section_row(label: str) -> dict[str, str]:
    return {"Characteristic": label, "eICU (N=777)": "", "MIMIC-IV (N=187)": ""}


@dataclass
class ArtifactSources:
    eicu: pd.DataFrame
    mimic: pd.DataFrame
    burden_summary: pd.DataFrame
    exclusion: pd.DataFrame
    flow_eicu: pd.DataFrame
    flow_mimic: pd.DataFrame
    v1_coef: pd.DataFrame
    v1_perf: pd.DataFrame
    v1_cal: pd.DataFrame
    v2_coef: pd.DataFrame
    v2_lr: pd.DataFrame
    v2_strat: pd.DataFrame
    v2_zero: pd.DataFrame
    v2_perf: pd.DataFrame
    rcs_lr: pd.DataFrame
    rcs_perf: pd.DataFrame
    rcs_curve: pd.DataFrame
    rcs_dist: pd.DataFrame
    winsor_coef: pd.DataFrame
    winsor_perf: pd.DataFrame
    winsor_lr: pd.DataFrame
    winsor_curve: pd.DataFrame
    hospital_coef: pd.DataFrame
    hospital_perf: pd.DataFrame
    map60_coef: pd.DataFrame
    map60_perf: pd.DataFrame
    vif: pd.DataFrame
    condition: pd.DataFrame
    spearman: pd.DataFrame
    qc_summary: pd.DataFrame
    missingness: pd.DataFrame
    rule_checks: pd.DataFrame


def load_sources() -> ArtifactSources:
    return ArtifactSources(
        eicu=read_csv_auto(source("03_outputs/tables/covariates/20260620_covariates_v2_baseline/eicu_analysis_covariates_v1_landmark_main.csv")),
        mimic=read_csv_auto(source("03_outputs/tables/covariates/20260620_covariates_v2_baseline/mimic_analysis_covariates_v1_landmark_main.csv")),
        burden_summary=read_csv_auto(source("03_outputs/tables/burden/20260620_burden24h_v2/burden_v2_summary.csv")),
        exclusion=read_csv_auto(source("03_outputs/tables/burden/20260620_burden24h_v2/burden_v2_exclusion_comparison.csv")),
        flow_eicu=read_csv_auto(source("03_outputs/cohort_flow/20260620_icuoutcome_fix/eicu_cohort_flow.csv")),
        flow_mimic=read_csv_auto(source("03_outputs/cohort_flow/20260620_icuoutcome_fix/mimic_cohort_flow.csv")),
        v1_coef=read_csv_auto(source("03_outputs/tables/models/20260620_adjusted_model_v1/adjusted_model_v1_coefficients.csv")),
        v1_perf=read_csv_auto(source("03_outputs/tables/models/20260620_adjusted_model_v1/adjusted_model_v1_performance.csv")),
        v1_cal=read_csv_auto(source("03_outputs/tables/models/20260620_adjusted_model_v1/adjusted_model_v1_calibration_groups.csv")),
        v2_coef=read_csv_auto(source("03_outputs/tables/models/20260620_adjusted_model_v2_interaction/adjusted_model_v2_coefficients.csv")),
        v2_lr=read_csv_auto(source("03_outputs/tables/models/20260620_adjusted_model_v2_interaction/adjusted_model_v2_lr_tests.csv")),
        v2_strat=read_csv_auto(source("03_outputs/tables/models/20260620_adjusted_model_v2_interaction/adjusted_model_v2_hypoxemia_stratified_hypotension_effect.csv")),
        v2_zero=read_csv_auto(source("03_outputs/tables/models/20260620_adjusted_model_v2_interaction/adjusted_model_v2_exposure_zero_summary.csv")),
        v2_perf=read_csv_auto(source("03_outputs/tables/models/20260620_adjusted_model_v2_interaction/adjusted_model_v2_performance.csv")),
        rcs_lr=read_csv_auto(source("03_outputs/tables/models/20260620_adjusted_model_v3b_hypotension_rcs_limited/adjusted_model_v3b_rcs_lr_tests.csv")),
        rcs_perf=read_csv_auto(source("03_outputs/tables/models/20260620_adjusted_model_v3b_hypotension_rcs_limited/adjusted_model_v3b_rcs_performance.csv")),
        rcs_curve=read_csv_auto(source("03_outputs/tables/models/20260620_adjusted_model_v3b_hypotension_rcs_limited/adjusted_model_v3b_rcs_prediction_curves.csv")),
        rcs_dist=read_csv_auto(source("03_outputs/tables/models/20260620_adjusted_model_v3b_hypotension_rcs_limited/adjusted_model_v3b_hypotension_distribution.csv")),
        winsor_coef=read_csv_auto(source("03_outputs/tables/sensitivity/20260621_analysis_closure_v1/winsor99_model_coefficients.csv")),
        winsor_perf=read_csv_auto(source("03_outputs/tables/sensitivity/20260621_analysis_closure_v1/winsor99_model_performance.csv")),
        winsor_lr=read_csv_auto(source("03_outputs/tables/sensitivity/20260621_analysis_closure_v1/winsor99_rcs_lr_tests.csv")),
        winsor_curve=read_csv_auto(source("03_outputs/tables/sensitivity/20260621_analysis_closure_v1/winsor99_rcs_prediction_curve.csv")),
        hospital_coef=read_csv_auto(source("03_outputs/tables/sensitivity/20260621_analysis_closure_v1/hospital_death_model_coefficients.csv")),
        hospital_perf=read_csv_auto(source("03_outputs/tables/sensitivity/20260621_analysis_closure_v1/hospital_death_model_performance.csv")),
        map60_coef=read_csv_auto(source("03_outputs/tables/sensitivity/20260621_analysis_closure_v1/map60_model_coefficients.csv")),
        map60_perf=read_csv_auto(source("03_outputs/tables/sensitivity/20260621_analysis_closure_v1/map60_model_performance.csv")),
        vif=read_csv_auto(source("03_outputs/tables/sensitivity/20260621_analysis_closure_v1/main_model_vif.csv")),
        condition=read_csv_auto(source("03_outputs/tables/sensitivity/20260621_analysis_closure_v1/main_model_condition_number.csv")),
        spearman=read_csv_auto(source("03_outputs/tables/sensitivity/20260621_analysis_closure_v1/spearman_correlations.csv")),
        qc_summary=read_csv_auto(source("03_outputs/tables/qc/20260620_final_dataset_qc_v1/final_dataset_qc_summary.csv")),
        missingness=read_csv_auto(source("03_outputs/tables/qc/20260620_final_dataset_qc_v1/final_dataset_missingness.csv")),
        rule_checks=read_csv_auto(source("03_outputs/tables/qc/20260620_final_dataset_qc_v1/final_dataset_rule_checks.csv")),
    )


def build_table1(src: ArtifactSources) -> pd.DataFrame:
    e, m = src.eicu, src.mimic
    e_surv, e_dead = e[e["death_icu"].eq(0)], e[e["death_icu"].eq(1)]
    m_surv, m_dead = m[m["death_icu"].eq(0)], m[m["death_icu"].eq(1)]
    columns = [
        "Characteristic",
        f"eICU Overall (N={len(e)})",
        f"eICU ICU survivor (N={len(e_surv)})",
        f"eICU ICU death (N={len(e_dead)})",
        "eICU P value",
        f"MIMIC-IV Overall (N={len(m)})",
        f"MIMIC-IV ICU survivor (N={len(m_surv)})",
        f"MIMIC-IV ICU death (N={len(m_dead)})",
        "MIMIC-IV P value",
    ]
    rows: list[dict[str, str]] = []

    def empty_row(label: str) -> dict[str, str]:
        row = {col: "" for col in columns}
        row["Characteristic"] = label
        return row

    def add_n(label: str) -> None:
        rows.append({
            columns[0]: label,
            columns[1]: fmt_int(len(e)),
            columns[2]: fmt_int(len(e_surv)),
            columns[3]: fmt_int(len(e_dead)),
            columns[4]: "",
            columns[5]: fmt_int(len(m)),
            columns[6]: fmt_int(len(m_surv)),
            columns[7]: fmt_int(len(m_dead)),
            columns[8]: "",
        })

    def add_cont(label: str, col: str, digits: int) -> None:
        rows.append({
            columns[0]: label,
            columns[1]: med_iqr(e[col], digits),
            columns[2]: med_iqr(e_surv[col], digits),
            columns[3]: med_iqr(e_dead[col], digits),
            columns[4]: fmt_p(mann_whitney_p(e[col], e["death_icu"])),
            columns[5]: med_iqr(m[col], digits),
            columns[6]: med_iqr(m_surv[col], digits),
            columns[7]: med_iqr(m_dead[col], digits),
            columns[8]: fmt_p(mann_whitney_p(m[col], m["death_icu"])),
        })

    def add_binary(label: str, col: str, sex_denominator: bool = False) -> None:
        e_over = n_pct(e[col]) if sex_denominator else n_pct(e[col], total=len(e))
        e_surv_val = n_pct(e_surv[col]) if sex_denominator else n_pct(e_surv[col], total=len(e_surv))
        e_dead_val = n_pct(e_dead[col]) if sex_denominator else n_pct(e_dead[col], total=len(e_dead))
        m_over = n_pct(m[col]) if sex_denominator else n_pct(m[col], total=len(m))
        m_surv_val = n_pct(m_surv[col]) if sex_denominator else n_pct(m_surv[col], total=len(m_surv))
        m_dead_val = n_pct(m_dead[col]) if sex_denominator else n_pct(m_dead[col], total=len(m_dead))
        rows.append({
            columns[0]: label,
            columns[1]: e_over,
            columns[2]: e_surv_val,
            columns[3]: e_dead_val,
            columns[4]: fmt_p(binary_p(e[col], e["death_icu"])),
            columns[5]: m_over,
            columns[6]: m_surv_val,
            columns[7]: m_dead_val,
            columns[8]: fmt_p(binary_p(m[col], m["death_icu"])),
        })

    def add_binary_descriptive(label: str, col: str) -> None:
        rows.append({
            columns[0]: label,
            columns[1]: n_pct(e[col], total=len(e)),
            columns[2]: n_pct(e_surv[col], total=len(e_surv)),
            columns[3]: n_pct(e_dead[col], total=len(e_dead)),
            columns[4]: "",
            columns[5]: n_pct(m[col], total=len(m)),
            columns[6]: n_pct(m_surv[col], total=len(m_surv)),
            columns[7]: n_pct(m_dead[col], total=len(m_dead)),
            columns[8]: "",
        })

    e = e.copy()
    m = m.copy()
    e["zero_hypotension_burden"] = e["hypotension_twa"].eq(0).astype(int)
    e["zero_hypoxemia_burden"] = e["hypoxemia_twa"].eq(0).astype(int)
    m["zero_hypotension_burden"] = m["hypotension_twa"].eq(0).astype(int)
    m["zero_hypoxemia_burden"] = m["hypoxemia_twa"].eq(0).astype(int)
    e_surv, e_dead = e[e["death_icu"].eq(0)], e[e["death_icu"].eq(1)]
    m_surv, m_dead = m[m["death_icu"].eq(0)], m[m["death_icu"].eq(1)]

    rows.append(empty_row("Population and outcomes"))
    add_n("ICU stays, n")
    add_cont("Age, years, median [IQR]", "age", 1)
    add_binary("Male sex, n (%)", "sex_male", sex_denominator=True)

    rows.append(empty_row("Neurologic severity"))
    for label, col in [
        ("GCS total, median [IQR]", "gcs_total"),
        ("GCS eye component, median [IQR]", "gcs_eyes"),
        ("GCS motor component, median [IQR]", "gcs_motor"),
        ("GCS verbal component, median [IQR]", "gcs_verbal"),
    ]:
        add_cont(label, col, 0)

    rows.append(empty_row("Early physiologic monitoring and burden"))
    add_cont("MAP effective observation, h, median [IQR]", "map_effective_hours", 1)
    add_cont("SpO2 effective observation, h, median [IQR]", "spo2_effective_hours", 1)
    add_cont("Hypotension burden (MAP <65), median [IQR]", "hypotension_twa", 3)
    add_cont("Hypoxemia burden (SpO2 <90), median [IQR]", "hypoxemia_twa", 3)
    add_binary("Zero hypotension burden, n (%)", "zero_hypotension_burden")
    add_binary("Zero hypoxemia burden, n (%)", "zero_hypoxemia_burden")

    rows.append(empty_row("TBI subtype descriptors"))
    subtype_labels = [
        ("Traumatic subdural hematoma, n (%)", "tbi_subdural"),
        ("Traumatic subarachnoid hemorrhage, n (%)", "tbi_subarachnoid"),
        ("Epidural hematoma, n (%)", "tbi_epidural"),
        ("Intracerebral hemorrhage, n (%)", "tbi_intracerebral_hemorrhage"),
        ("Contusion or laceration, n (%)", "tbi_contusion_or_laceration"),
        ("Diffuse axonal injury, n (%)", "tbi_diffuse_axonal_injury"),
        ("Herniation, n (%)", "tbi_herniation"),
        ("Cerebral edema, n (%)", "tbi_edema"),
    ]
    for label, col in subtype_labels:
        add_binary(label, col)

    rows.append(empty_row("First-24h treatment descriptors"))
    add_binary("Mechanical ventilation within 24 h, n (%)", "mechanical_vent_24h")
    add_binary("Vasopressor use within 24 h, n (%)", "vasopressor_24h")

    rows.append(empty_row("Comorbidities"))
    comorbidity_labels = [
        ("Congestive heart failure, n (%)", "hx_congestive_heart_failure"),
        ("Diabetes, n (%)", "hx_diabetes"),
        ("Chronic pulmonary disease, n (%)", "hx_chronic_pulmonary_disease"),
        ("Renal disease, n (%)", "hx_renal_disease"),
        ("Liver disease, n (%)", "hx_liver_disease"),
        ("Malignancy, n (%)", "hx_malignancy"),
        ("Cerebrovascular disease, n (%)", "hx_cerebrovascular_disease"),
    ]
    for label, col in comorbidity_labels:
        add_binary(label, col)

    rows.append(empty_row("Outcome descriptors"))
    add_binary_descriptive("Hospital death, n (%)", "death_hospital")

    return pd.DataFrame(rows)


def coeff(df: pd.DataFrame, variable: str, model: str | None = None) -> pd.Series:
    query = df[df["variable"].eq(variable)]
    if model is not None and "model" in query.columns:
        query = query[query["model"].eq(model)]
    if query.empty:
        raise KeyError(f"Missing coefficient for {variable}, model={model}")
    return query.iloc[0]


def perf(df: pd.DataFrame, dataset: str = "MIMIC_IV_external_validation", model: str | None = None) -> pd.Series:
    query = df[df["dataset"].eq(dataset)]
    if model is not None and "model" in query.columns:
        query = query[query["model"].eq(model)]
    if query.empty:
        raise KeyError(f"Missing performance for {dataset}, model={model}")
    return query.iloc[0]


def build_table2(src: ArtifactSources) -> pd.DataFrame:
    rows: list[dict[str, str]] = []

    def add(section: str, analysis: str, effect: pd.Series, p: float, perf_row: pd.Series, eicu_n_events: str, note: str = "") -> None:
        rows.append({
            "Section": section,
            "Analysis": analysis,
            "eICU analysis set": eicu_n_events,
            "Effect estimate": fmt_or_ci(effect),
            "P value": fmt_p(p),
            "MIMIC-IV C-index": fmt_float(perf_row["auc_c_index"], 3),
            "MIMIC-IV calibration slope": fmt_float(perf_row["calibration_slope_logistic"], 3),
            "MIMIC-IV calibration intercept": fmt_float(perf_row["calibration_intercept_logistic"], 3),
            "Note": note,
        })

    main_effect = coeff(src.v1_coef, "hypotension_twa_per_sd")
    main_perf = perf(src.v1_perf)
    eicu_main = perf(src.v1_perf, dataset="eICU_derivation")
    add(
        "Primary",
        "ICU death; MAP <65 hypotension burden, per eICU SD",
        main_effect,
        main_effect["p_value"],
        main_perf,
        f"{fmt_int(eicu_main['n'])}/{fmt_int(eicu_main['events'])}",
        "Locked main-effect logistic model.",
    )

    hosp_effect = coeff(src.hospital_coef, "hypotension_twa_per_sd")
    hosp_perf = perf(src.hospital_perf)
    eicu_hosp = perf(src.hospital_perf, dataset="eICU_derivation")
    add(
        "Sensitivity",
        "Hospital death; MAP <65 hypotension burden, per eICU SD",
        hosp_effect,
        hosp_effect["p_value"],
        hosp_perf,
        f"{fmt_int(eicu_hosp['n'])}/{fmt_int(eicu_hosp['events'])}",
        "762 = 777 landmark stays minus 13 missing hospital outcomes and 2 sex-missing stays.",
    )

    map60_effect = coeff(src.map60_coef, "hypotension_twa_per_sd")
    map60_perf = perf(src.map60_perf)
    eicu_map60 = perf(src.map60_perf, dataset="eICU_derivation")
    add(
        "Sensitivity",
        "ICU death; MAP <60 hypotension burden, per eICU SD",
        map60_effect,
        map60_effect["p_value"],
        map60_perf,
        f"{fmt_int(eicu_map60['n'])}/{fmt_int(eicu_map60['events'])}",
        "MAP<60 burden recalculated; MIMIC-IV validation N remained 187/26.",
    )

    winsor_effect = coeff(src.winsor_coef[src.winsor_coef["model"].eq("winsor99_linear")], "hypotension_twa_per_sd")
    winsor_perf = perf(src.winsor_perf, model="winsor99_linear")
    eicu_winsor = perf(src.winsor_perf, dataset="eICU_derivation", model="winsor99_linear")
    cap = src.winsor_coef["cap_value_eicu_p99"].dropna().iloc[0]
    add(
        "Sensitivity",
        "ICU death; 99th-percentile winsorized MAP <65 burden",
        winsor_effect,
        winsor_effect["p_value"],
        winsor_perf,
        f"{fmt_int(eicu_winsor['n'])}/{fmt_int(eicu_winsor['events'])}",
        f"eICU P99 cap = {cap:.4f}.",
    )

    hypox_effect = coeff(src.v1_coef, "hypoxemia_twa_per_sd")
    add(
        "Secondary hypoxemia",
        "Hypoxemia burden main effect, per eICU SD",
        hypox_effect,
        hypox_effect["p_value"],
        main_perf,
        f"{fmt_int(eicu_main['n'])}/{fmt_int(eicu_main['events'])}",
        "Prespecified common exposure; sparse zero-inflated burden.",
    )

    return pd.DataFrame(rows)


def build_supplementary_tables(src: ArtifactSources) -> dict[str, pd.DataFrame]:
    supp: dict[str, pd.DataFrame] = {}

    flow_rows = []
    for df in [src.flow_eicu, src.flow_mimic]:
        for _, row in df[df["section"].eq("cohort_flow")].iterrows():
            flow_rows.append({
                "Database": row["database_name"],
                "Step": row["metric"],
                "N": fmt_int(row["n"]),
                "ICU deaths": fmt_int(row["deaths"]),
                "ICU mortality, %": fmt_float(row["mortality_pct"], 1),
                "Notes": row.get("notes", ""),
            })
    supp["Table_S1_Cohort_Flow_Source_Counts"] = pd.DataFrame(flow_rows)

    excl_rows = src.exclusion.copy()
    excl_rows["ICU deaths"] = excl_rows.apply(lambda r: n_pct_from_counts(r["icu_deaths_n"], r["n"], 1), axis=1)
    excl_rows["Hospital deaths"] = excl_rows.apply(lambda r: n_pct_from_counts(r["hospital_deaths_n"], r["n"], 1), axis=1)
    supp["Table_S2_Landmark_And_Coverage_Exclusions"] = excl_rows[[
        "database_name", "section", "group", "n", "ICU deaths", "Hospital deaths",
        "gcs_median", "window_hours_median", "hypotension_twa_median", "hypoxemia_twa_median",
    ]].rename(columns={
        "database_name": "Database",
        "section": "Gate",
        "group": "Group",
        "n": "N",
        "gcs_median": "GCS median",
        "window_hours_median": "Window hours, median",
        "hypotension_twa_median": "Hypotension burden, median",
        "hypoxemia_twa_median": "Hypoxemia burden, median",
    })

    qc_rows = src.qc_summary.copy()
    supp["Table_S3_QC_And_Missingness_Summary"] = qc_rows[[
        "database_name", "n", "icu_deaths_n", "icu_death_pct", "hospital_deaths_n",
        "hospital_death_pct", "main_model_complete_case_n", "age_median",
        "gcs_total_median", "hypotension_twa_median", "hypoxemia_twa_median",
        "map_effective_hours_median", "spo2_effective_hours_median",
    ]].rename(columns={
        "database_name": "Database",
        "n": "N",
        "icu_deaths_n": "ICU deaths",
        "icu_death_pct": "ICU death, %",
        "hospital_deaths_n": "Hospital deaths",
        "hospital_death_pct": "Hospital death, %",
        "main_model_complete_case_n": "Main model complete-case N",
        "age_median": "Age median",
        "gcs_total_median": "GCS median",
        "hypotension_twa_median": "Hypotension burden median",
        "hypoxemia_twa_median": "Hypoxemia burden median",
        "map_effective_hours_median": "MAP observed hours median",
        "spo2_effective_hours_median": "SpO2 observed hours median",
    })

    zero = src.v2_zero.copy()
    supp["Table_S4_Exposure_Zero_Inflation"] = zero.rename(columns={
        "dataset": "Dataset",
        "variable": "Variable",
        "n": "N",
        "zero_n": "Zero N",
        "zero_pct": "Zero %",
        "nonzero_n": "Nonzero N",
        "nonzero_pct": "Nonzero %",
    })

    interaction_rows = []
    for _, row in src.v2_lr.iterrows():
        interaction_rows.append({
            "Analysis": row["comparison"],
            "Effect estimate": "LR test",
            "OR (95% CI)": "-",
            "P value": fmt_p(row["p_value"]),
            "N/events": "-",
            "C-index": "-",
        })
    interaction_terms = [
        (
            "Continuous hypotension x hypoxemia interaction",
            coeff(src.v2_coef, "hypotension_x_hypoxemia_z", model="continuous_interaction"),
            perf(src.v2_perf, dataset="eICU_derivation", model="continuous_interaction"),
            perf(src.v2_perf, model="continuous_interaction"),
        ),
        (
            "Hypotension x any hypoxemia-burden interaction",
            coeff(src.v2_coef, "hypotension_z_x_hypoxemia_any", model="hypoxemia_binary_interaction"),
            perf(src.v2_perf, dataset="eICU_derivation", model="hypoxemia_binary_interaction"),
            perf(src.v2_perf, model="hypoxemia_binary_interaction"),
        ),
    ]
    for label, coef_row, eicu_row, mimic_row in interaction_terms:
        interaction_rows.append({
            "Analysis": label,
            "Effect estimate": coef_row["label"],
            "OR (95% CI)": fmt_or_ci(coef_row),
            "P value": fmt_p(coef_row["p_value"]),
            "N/events": f"{fmt_int(eicu_row['n'])}/{fmt_int(eicu_row['events'])}",
            "C-index": f"eICU {fmt_float(eicu_row['auc_c_index'], 3)}; MIMIC-IV {fmt_float(mimic_row['auc_c_index'], 3)}",
        })
    for _, row in src.v2_strat.iterrows():
        interaction_rows.append({
            "Analysis": f"Hypotension effect in {row['stratum']}",
            "Effect estimate": row["label"],
            "OR (95% CI)": fmt_or_ci(row),
            "P value": fmt_p(row["p_value"]),
            "N/events": f"{fmt_int(row['n'])}/{fmt_int(row['events'])}",
            "C-index": fmt_float(row["auc_c_index"], 3),
        })
    supp["Table_S5_Hypoxemia_Interaction_And_Stratified_Analyses"] = pd.DataFrame(interaction_rows)

    rcs_rows = []
    for _, row in src.rcs_lr.iterrows():
        rcs_rows.append({
            "Model": row["model"],
            "Comparison": row["comparison"],
            "LR chi-square": fmt_float(row["lr_chisq"], 3),
            "df": fmt_int(row["df"]),
            "P value": fmt_p(row["p_value"]),
        })
    for _, row in src.winsor_lr.iterrows():
        rcs_rows.append({
            "Model": row["model"],
            "Comparison": row["comparison"],
            "LR chi-square": fmt_float(row["lr_chisq"], 3),
            "df": fmt_int(row["df"]),
            "P value": fmt_p(row["p_value"]),
        })
    supp["Table_S6_RCS_And_Winsorized_RCS_Tests"] = pd.DataFrame(rcs_rows)

    performance = pd.concat([
        src.v1_perf.assign(model="linear_hypotension_v1"),
        src.rcs_perf,
        src.winsor_perf,
        src.hospital_perf,
        src.map60_perf,
    ], ignore_index=True, sort=False)
    supp["Table_S7_Discrimination_And_Calibration"] = performance[[
        "model", "dataset", "n", "events", "event_rate", "auc_c_index", "brier_score",
        "mean_predicted_risk", "observed_risk", "calibration_intercept_logistic",
        "calibration_slope_logistic",
    ]].rename(columns={
        "model": "Model",
        "dataset": "Dataset",
        "n": "N",
        "events": "Events",
        "event_rate": "Event rate",
        "auc_c_index": "C-index",
        "brier_score": "Brier score",
        "mean_predicted_risk": "Mean predicted risk",
        "observed_risk": "Observed risk",
        "calibration_intercept_logistic": "Calibration intercept",
        "calibration_slope_logistic": "Calibration slope",
    })

    vif_cond = src.vif.copy()
    condition = src.condition.iloc[0]
    vif_cond["condition_number_standardized_design"] = condition["condition_number_standardized_design"]
    supp["Table_S8_Collinearity_Diagnostics"] = vif_cond.rename(columns={
        "variable": "Variable",
        "r_squared_on_other_predictors": "R2 on other predictors",
        "tolerance": "Tolerance",
        "vif": "VIF",
    })

    spearman = src.spearman[src.spearman["row_variable"].ne(src.spearman["column_variable"])].copy()
    supp["Table_S9_Spearman_Correlations"] = spearman.rename(columns={
        "dataset": "Dataset",
        "row_variable": "Row variable",
        "column_variable": "Column variable",
        "spearman_rho": "Spearman rho",
    })

    return supp


def write_markdown_table(path: Path, title: str, df: pd.DataFrame, note: str = "") -> None:
    def clean(value: object) -> str:
        return str(value).replace("\n", "<br>").replace("|", "\\|")

    headers = [clean(c) for c in df.columns]
    body = [[clean(v) for v in row] for row in df.astype(str).to_numpy()]
    lines = [
        f"# {title}",
        "",
        "| " + " | ".join(headers) + " |",
        "| " + " | ".join(["---"] * len(headers)) + " |",
    ]
    lines.extend("| " + " | ".join(row) + " |" for row in body)
    if note:
        lines += ["", f"Note: {note}"]
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def write_table1_markdown(path: Path, title: str, df: pd.DataFrame, note: str = "") -> None:
    lines = [
        f"# {title}",
        "",
        "<table>",
        "  <thead>",
        "    <tr>",
        "      <th rowspan=\"2\">Characteristic</th>",
        "      <th colspan=\"4\">eICU</th>",
        "      <th colspan=\"4\">MIMIC-IV</th>",
        "    </tr>",
        "    <tr>",
        "      <th>Overall</th><th>ICU survivor</th><th>ICU death</th><th>P value</th>",
        "      <th>Overall</th><th>ICU survivor</th><th>ICU death</th><th>P value</th>",
        "    </tr>",
        "  </thead>",
        "  <tbody>",
    ]
    for _, row in df.iterrows():
        values = ["" if pd.isna(v) else str(v) for v in row.to_list()]
        is_section = values[0].strip() and all(v.strip() == "" for v in values[1:])
        if is_section:
            lines.append(f"    <tr><td colspan=\"9\"><strong>{escape(values[0])}</strong></td></tr>")
        else:
            cells = "".join(f"<td>{escape(v)}</td>" for v in values)
            lines.append(f"    <tr>{cells}</tr>")
    lines += ["  </tbody>", "</table>"]
    if note:
        lines += ["", f"Note: {note}"]
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def set_cell_text(cell, text: str, bold: bool = False, font_size: float = 10.5) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = 0
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run("" if text is None else str(text))
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(font_size)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill.replace("#", ""))
    tc_pr.append(shd)


def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = "w:{}".format(edge)
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key in ["sz", "val", "color", "space"]:
                if key in edge_data:
                    element.set(qn(f"w:{key}"), str(edge_data[key]))


def set_table_borders_minimal(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(
                cell,
                top={"val": "nil"},
                left={"val": "nil"},
                right={"val": "nil"},
                bottom={"val": "nil"},
            )
    for cell in table.rows[0].cells:
        set_cell_border(
            cell,
            top={"val": "single", "sz": "12", "color": "000000"},
            bottom={"val": "single", "sz": "8", "color": "000000"},
            left={"val": "nil"},
            right={"val": "nil"},
        )
    for cell in table.rows[-1].cells:
        set_cell_border(
            cell,
            bottom={"val": "single", "sz": "12", "color": "000000"},
            left={"val": "nil"},
            right={"val": "nil"},
        )


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tbl_cell_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_column_widths(table, widths_in: list[float]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths_in):
            if idx >= len(row.cells):
                continue
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def add_docx_table(doc, title: str, df: pd.DataFrame, note: str, col_widths: list[float] | None = None) -> None:
    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(6)
    heading.paragraph_format.space_after = Pt(4)
    run = heading.add_run(title)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.autofit = col_widths is None
    set_table_cell_margins(table)
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, col in enumerate(df.columns):
        set_cell_text(header.cells[idx], str(col), bold=True, font_size=9.5)
        shade_cell(header.cells[idx], "F2F4F7")

    for _, data_row in df.iterrows():
        row = table.add_row()
        is_section = str(data_row.iloc[0]).strip() and all(str(x).strip() == "" for x in data_row.iloc[1:])
        for idx, val in enumerate(data_row):
            set_cell_text(row.cells[idx], str(val), bold=is_section, font_size=9.2 if len(df.columns) > 5 else 10.0)
            if is_section:
                shade_cell(row.cells[idx], "E8EEF5")

    if col_widths is not None:
        set_column_widths(table, col_widths)
    set_table_borders_minimal(table)

    note_p = doc.add_paragraph()
    note_p.paragraph_format.space_before = Pt(4)
    note_p.paragraph_format.space_after = Pt(8)
    note_run = note_p.add_run("Note: " + note)
    note_run.font.name = "Times New Roman"
    note_run.font.size = Pt(8.5)


def add_docx_mixed_table1(doc, df: pd.DataFrame, note: str) -> None:
    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(6)
    heading.paragraph_format.space_after = Pt(4)
    run = heading.add_run("Mixed baseline characteristics")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)

    table = doc.add_table(rows=2, cols=len(df.columns))
    table.autofit = False
    set_table_cell_margins(table, top=60, bottom=60, start=80, end=80)
    set_repeat_table_header(table.rows[0])
    set_repeat_table_header(table.rows[1])

    top = table.rows[0].cells
    bottom = table.rows[1].cells
    top[1].merge(top[4])
    top[5].merge(top[8])

    set_cell_text(table.rows[0].cells[0], "Characteristic", bold=True, font_size=8.5)
    set_cell_text(table.rows[1].cells[0], "", bold=True, font_size=8.0)
    set_cell_text(table.rows[0].cells[1], "eICU", bold=True, font_size=8.5)
    set_cell_text(table.rows[0].cells[5], "MIMIC-IV", bold=True, font_size=8.5)
    for idx, label in enumerate(["Overall", "ICU survivor", "ICU death", "P value", "Overall", "ICU survivor", "ICU death", "P value"], start=1):
        set_cell_text(table.rows[1].cells[idx], label, bold=True, font_size=8.0)

    for _, data_row in df.iterrows():
        row = table.add_row()
        is_section = str(data_row.iloc[0]).strip() and all(str(x).strip() == "" for x in data_row.iloc[1:])
        for idx, val in enumerate(data_row):
            set_cell_text(row.cells[idx], str(val), bold=is_section, font_size=7.4 if not is_section else 8.0)

    set_column_widths(table, [2.45, 1.04, 1.04, 1.04, 0.55, 1.04, 1.04, 1.04, 0.55])
    set_table_borders_minimal(table)
    for cell in table.rows[1].cells:
        set_cell_border(
            cell,
            bottom={"val": "single", "sz": "8", "color": "000000"},
            left={"val": "nil"},
            right={"val": "nil"},
        )

    note_p = doc.add_paragraph()
    note_p.paragraph_format.space_before = Pt(4)
    note_p.paragraph_format.space_after = Pt(8)
    note_run = note_p.add_run("Note: " + note)
    note_run.font.name = "Times New Roman"
    note_run.font.size = Pt(8.5)


def write_docx_tables(table1: pd.DataFrame, table2: pd.DataFrame, supp: dict[str, pd.DataFrame]) -> list[Path]:
    outputs: list[Path] = []

    def save_docx(doc, path: Path) -> Path:
        try:
            doc.save(path)
            return path
        except PermissionError:
            fallback = path.with_name(f"{path.stem}_no_fill{path.suffix}") if "Table_1" in path.name else path.with_name(f"{path.stem}_updated{path.suffix}")
            doc.save(fallback)
            return fallback

    def setup_doc(title: str, landscape: bool = False):
        doc = Document()
        section = doc.sections[0]
        if landscape:
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = Inches(11)
            section.page_height = Inches(8.5)
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
        else:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        styles = doc.styles
        styles["Normal"].font.name = "Times New Roman"
        styles["Normal"].font.size = Pt(10.5)
        title_p = doc.add_paragraph()
        title_p.paragraph_format.space_after = Pt(8)
        title_run = title_p.add_run(title)
        title_run.bold = True
        title_run.font.name = "Times New Roman"
        title_run.font.size = Pt(12)
        return doc

    doc1 = setup_doc("Table 1. Mixed baseline characteristics of the landmark severe TBI cohorts", landscape=True)
    add_docx_mixed_table1(
        doc1,
        table1,
        "Continuous variables are summarized as median [IQR]; categorical variables as n (%). P values compare ICU survivors and ICU deaths within each database using Mann-Whitney U tests for continuous variables and Fisher exact or chi-square tests for binary variables. No P values compare eICU with MIMIC-IV because the databases served derivation and validation roles. Hospital death is shown last as a descriptive outcome descriptor; no P value is reported for that row because ICU death defines the stratification. First-24h mechanical ventilation and vasopressor use are descriptive treatment-window variables and were not included in the primary model because they overlap the exposure window and may represent treatment response or mediation. MIMIC-IV herniation counts reflect source coding and should not be interpreted as true absence or as a cross-database subtype contrast. Sex percentages use non-missing denominators; sex was missing for 2 eICU survivors.",
    )
    path1 = TABLE_DIR / "Table_1_Baseline_Characteristics.docx"
    outputs.append(save_docx(doc1, path1))

    doc2 = setup_doc("Table 2. Primary model and prespecified sensitivity analyses", landscape=True)
    add_docx_table(
        doc2,
        "Main effect and sensitivity analyses",
        table2,
        "Effect estimates are odds ratios per eICU SD unless otherwise stated. Models were fit in eICU and externally evaluated in MIMIC-IV using locked eICU coefficients. RCS analyses are reported in the supplement because nonlinearity attenuated after 99th-percentile winsorization.",
        [0.95, 2.05, 0.85, 1.15, 0.6, 0.7, 0.82, 0.82, 1.35],
    )
    path2 = TABLE_DIR / "Table_2_Main_Model_And_Sensitivity_Analyses.docx"
    outputs.append(save_docx(doc2, path2))

    doc_s = setup_doc("Supplementary tables", landscape=True)
    for idx, (name, df) in enumerate(supp.items(), start=1):
        short = name.replace("_", " ")
        add_docx_table(
            doc_s,
            f"Supplementary {short}",
            df,
            "Generated from locked project outputs. Values should be interpreted as descriptive or sensitivity analyses where indicated.",
            None,
        )
        if idx != len(supp):
            doc_s.add_page_break()
    path_s = TABLE_DIR / "Supplementary_Tables.docx"
    outputs.append(save_docx(doc_s, path_s))
    return outputs


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    font_name = "arialbd.ttf" if bold else "arial.ttf"
    path = Path("C:/Windows/Fonts") / font_name
    if path.exists():
        return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def text_width(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.ImageFont) -> float:
    if text == "":
        return 0
    box = draw.textbbox((0, 0), text, font=fnt)
    return box[2] - box[0]


def wrap_pixels(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.ImageFont, max_width: int) -> list[str]:
    lines: list[str] = []
    for raw in str(text).split("\n"):
        words = raw.split()
        if not words:
            lines.append("")
            continue
        current = words[0]
        for word in words[1:]:
            candidate = f"{current} {word}"
            if text_width(draw, candidate, fnt) <= max_width:
                current = candidate
            else:
                lines.append(current)
                current = word
        lines.append(current)
    return lines


def draw_multiline(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    fnt: ImageFont.ImageFont,
    fill: str = "#111111",
    align: str = "center",
    valign: str = "center",
    spacing: float = 1.14,
) -> None:
    x0, y0, x1, y1 = box
    lines = wrap_pixels(draw, text, fnt, max(1, x1 - x0 - 24))
    line_h = max(1, int(fnt.size * spacing)) if hasattr(fnt, "size") else 16
    total_h = line_h * len(lines)
    if valign == "center":
        y = y0 + (y1 - y0 - total_h) / 2
    elif valign == "bottom":
        y = y1 - total_h
    else:
        y = y0
    for line in lines:
        w = text_width(draw, line, fnt)
        if align == "center":
            x = x0 + (x1 - x0 - w) / 2
        elif align == "right":
            x = x1 - w
        else:
            x = x0
        draw.text((x, y), line, font=fnt, fill=fill)
        y += line_h


def draw_round_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    fill: str,
    outline: str,
    fnt: ImageFont.ImageFont,
    text_fill: str = "#111111",
    width: int = 5,
    radius: int = 34,
) -> None:
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)
    draw_multiline(draw, box, text, fnt, fill=text_fill)


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], fill: str = "#6B7280", width: int = 5) -> None:
    draw.line([start, end], fill=fill, width=width)
    x0, y0 = start
    x1, y1 = end
    angle = math.atan2(y1 - y0, x1 - x0)
    length = 22
    spread = math.pi / 6
    p1 = (x1 - length * math.cos(angle - spread), y1 - length * math.sin(angle - spread))
    p2 = (x1 - length * math.cos(angle + spread), y1 - length * math.sin(angle + spread))
    draw.polygon([(x1, y1), p1, p2], fill=fill)


def svg_text_lines(x: float, y: float, lines: list[str], size: int, fill: str = "#111111", weight: str = "normal", anchor: str = "middle", line_height: float = 1.18) -> str:
    tspans = []
    for idx, line in enumerate(lines):
        dy = 0 if idx == 0 else size * line_height
        tspans.append(f'<tspan x="{x:.1f}" dy="{dy:.1f}">{escape(line)}</tspan>')
    return f'<text x="{x:.1f}" y="{y:.1f}" text-anchor="{anchor}" font-family="Arial" font-size="{size}" font-weight="{weight}" fill="{fill}">{"".join(tspans)}</text>'


def svg_wrapped_text(x: float, y: float, text: str, size: int, max_chars: int, **kwargs) -> str:
    lines: list[str] = []
    for part in str(text).split("\n"):
        wrapped = textwrap.wrap(part, width=max_chars, break_long_words=False) or [""]
        lines.extend(wrapped)
    return svg_text_lines(x, y, lines, size, **kwargs)


def svg_box(x: float, y: float, w: float, h: float, text: str, fill: str, stroke: str, size: int = 30, weight: str = "normal", max_chars: int = 26) -> str:
    lines: list[str] = []
    for part in str(text).split("\n"):
        lines.extend(textwrap.wrap(part, width=max_chars, break_long_words=False) or [""])
    start_y = y + h / 2 - ((len(lines) - 1) * size * 1.16) / 2
    return (
        f'<rect x="{x:.1f}" y="{y:.1f}" width="{w:.1f}" height="{h:.1f}" rx="34" fill="{fill}" stroke="{stroke}" stroke-width="5"/>'
        + svg_text_lines(x + w / 2, start_y, lines, size, weight=weight)
    )


def save_svg_png_tiff(img: Image.Image, svg: str, stem: str, dpi: int = 600) -> dict[str, str]:
    svg_path = FIGURE_DIR / f"{stem}.svg"
    png_path = FIGURE_DIR / f"{stem}_{dpi}dpi.png"
    tiff_path = FIGURE_DIR / f"{stem}_{dpi}dpi.tiff"
    svg_path.write_text(svg, encoding="utf-8")
    img.save(png_path, dpi=(dpi, dpi))
    img.save(tiff_path, dpi=(dpi, dpi), compression="tiff_lzw")
    return {"svg": str(svg_path), "png": str(png_path), "tiff": str(tiff_path), "pixel_size": f"{img.width}x{img.height}", "dpi": str(dpi)}


def build_figure1(src: ArtifactSources) -> dict[str, str]:
    W, H = 4200, 3000
    img = Image.new("RGB", (W, H), "white")
    draw = ImageDraw.Draw(img)
    draw_multiline(draw, (300, 80, W - 300, 150), "Study flow for severe TBI landmark cohorts", font(70, True), fill=COLORS["navy"])
    draw_multiline(draw, (300, 170, W - 300, 235), "Adults with strict intracranial TBI and GCS <=8; 24-h landmark with >=12 h MAP and SpO2 effective observation", font(34), fill=COLORS["gray"])

    columns = {
        "eICU derivation": {
            "x": 250,
            "edge": COLORS["green"],
            "light": COLORS["green_light"],
            "steps": [
                ("Adult first ICU unit stays", "N=165,795\nICU deaths=9,094"),
                ("Strict intracranial TBI diagnosis", "N=4,202\nICU deaths=327"),
                ("Valid APACHE GCS", "N=4,060\nICU deaths=321"),
                ("Severe TBI: GCS <=8", "N=989\nICU deaths=263"),
                ("24-h landmark survived/remained in ICU", "N=807\nExcluded early exit/death: 182"),
                (">=12 h MAP and SpO2 observation", "Final N=777\nICU deaths=158; hospital deaths=217"),
            ],
            "exclusion": "Survived/remained 24 h but failed coverage: N=30",
            "final": "Derivation model\ncomplete-case N=775",
        },
        "MIMIC-IV external validation": {
            "x": 2240,
            "edge": COLORS["blue"],
            "light": COLORS["blue_light"],
            "steps": [
                ("Adult first ICU stays", "N=85,242\nICU deaths=6,226"),
                ("Strict TBI diagnosis", "N=3,202\nICU deaths=262"),
                ("First-day GCS available", "N=3,193\nICU deaths=257"),
                ("Severe TBI: GCS <=8", "N=223\nICU deaths=39"),
                ("24-h landmark survived/remained in ICU", "N=193\nExcluded early exit/death: 30"),
                (">=12 h MAP and SpO2 observation", "Final N=187\nICU deaths=26; hospital deaths=49"),
            ],
            "exclusion": "Survived/remained 24 h but failed coverage: N=6",
            "final": "External validation\nlocked eICU coefficients",
        },
    }

    y_positions = [455, 800, 1145, 1490, 1910, 2350]
    w, h = 1710, 235
    svg_parts = [
        f'<svg xmlns="http://www.w3.org/2000/svg" width="{W}" height="{H}" viewBox="0 0 {W} {H}">',
        '<rect width="100%" height="100%" fill="white"/>',
        svg_text_lines(W / 2, 120, ["Study flow for severe TBI landmark cohorts"], 70, fill=COLORS["navy"], weight="bold"),
        svg_wrapped_text(W / 2, 205, "Adults with strict intracranial TBI and GCS <=8; 24-h landmark with >=12 h MAP and SpO2 effective observation", 34, 130, fill=COLORS["gray"]),
    ]
    for title, cfg in columns.items():
        x = cfg["x"]
        draw_multiline(draw, (x, 300, x + w, 375), title, font(50, True), fill=cfg["edge"])
        svg_parts.append(svg_text_lines(x + w / 2, 350, [title], 50, fill=cfg["edge"], weight="bold"))
        for idx, ((label, value), y) in enumerate(zip(cfg["steps"], y_positions)):
            fc = cfg["light"]
            ec = cfg["edge"]
            if idx == 4:
                fc, ec = COLORS["purple_light"], COLORS["purple"]
            if idx == 5:
                fc, ec = COLORS["gold_light"], COLORS["gold"]
            box_text = f"{label}\n{value}"
            draw_round_box(draw, (x, y, x + w, y + h), box_text, fc, ec, font(34, idx == 5))
            svg_parts.append(svg_box(x, y, w, h, box_text, fc, ec, size=34, weight="bold" if idx == 5 else "normal", max_chars=36))
            if idx < len(y_positions) - 1:
                start = (x + w // 2, y + h + 24)
                end = (x + w // 2, y_positions[idx + 1] - 24)
                draw_arrow(draw, start, end)
                svg_parts.append(f'<line x1="{start[0]}" y1="{start[1]}" x2="{end[0]}" y2="{end[1]}" stroke="#6B7280" stroke-width="5" marker-end="url(#arrow)"/>')
        draw_round_box(draw, (x, 2700, x + w, 2858), cfg["final"], COLORS["gray_light"], cfg["edge"], font(34, True))
        draw_multiline(draw, (x, 2595, x + w, 2660), cfg["exclusion"], font(30), fill=COLORS["gray"])
        svg_parts.append(svg_box(x, 2700, w, 158, cfg["final"], COLORS["gray_light"], cfg["edge"], size=34, weight="bold", max_chars=35))
        svg_parts.append(svg_wrapped_text(x + w / 2, 2635, cfg["exclusion"], 30, 58, fill=COLORS["gray"]))

    foot = "MAP: mean arterial pressure; SpO2: peripheral oxygen saturation; TBI: traumatic brain injury. Early exit/death before 24 h is described but excluded by landmark design."
    draw_multiline(draw, (300, 2910, W - 300, 2980), foot, font(28), fill=COLORS["gray"])
    svg_parts.insert(1, '<defs><marker id="arrow" markerWidth="12" markerHeight="12" refX="10" refY="6" orient="auto"><path d="M0,0 L12,6 L0,12 z" fill="#6B7280"/></marker></defs>')
    svg_parts.append(svg_wrapped_text(W / 2, 2945, foot, 28, 145, fill=COLORS["gray"]))
    svg_parts.append("</svg>")
    return save_svg_png_tiff(img, "\n".join(svg_parts), "Figure_1_Study_Flow", dpi=600)


def build_figure2(src: ArtifactSources) -> dict[str, str]:
    W, H = 5200, 3000
    img = Image.new("RGB", (W, H), "white")
    draw = ImageDraw.Draw(img)
    draw_multiline(draw, (300, 70, W - 300, 165), "Main effect model and MIMIC-IV external validation", font(70, True), fill=COLORS["navy"])
    svg_parts = [
        f'<svg xmlns="http://www.w3.org/2000/svg" width="{W}" height="{H}" viewBox="0 0 {W} {H}">',
        '<rect width="100%" height="100%" fill="white"/>',
        svg_text_lines(W / 2, 125, ["Main effect model and MIMIC-IV external validation"], 70, fill=COLORS["navy"], weight="bold"),
    ]
    rows = [
        ("Primary ICU death\nMAP <65", coeff(src.v1_coef, "hypotension_twa_per_sd"), COLORS["navy"]),
        ("Hospital death\nalternative outcome", coeff(src.hospital_coef, "hypotension_twa_per_sd"), COLORS["green"]),
        ("MAP <60\nalternative threshold", coeff(src.map60_coef, "hypotension_twa_per_sd"), COLORS["blue"]),
        ("99% winsorized\nMAP <65", coeff(src.winsor_coef[src.winsor_coef["model"].eq("winsor99_linear")], "hypotension_twa_per_sd"), COLORS["gold"]),
    ]
    draw.text((260, 310), "A. Low blood pressure burden: main and sensitivity effects", font=font(44, True), fill=COLORS["navy"])
    svg_parts.append(svg_text_lines(260, 350, ["A. Low blood pressure burden: main and sensitivity effects"], 44, fill=COLORS["navy"], weight="bold", anchor="start"))
    plot = (1180, 620, 2390, 2200)
    label_x = 260
    estimate_x = 2470
    min_or, max_or = 0.85, 2.35

    def x_or(value: float) -> int:
        return int(plot[0] + (value - min_or) / (max_or - min_or) * (plot[2] - plot[0]))

    for tick in [1.0, 1.25, 1.5, 1.75, 2.0, 2.25]:
        x = x_or(tick)
        draw.line([(x, plot[1]), (x, plot[3])], fill="#E5E7EB", width=3)
        draw.text((x - 28, plot[3] + 24), f"{tick:.2g}", font=font(28), fill=COLORS["gray"])
        svg_parts.append(f'<line x1="{x}" y1="{plot[1]}" x2="{x}" y2="{plot[3]}" stroke="#E5E7EB" stroke-width="3"/>')
        svg_parts.append(svg_text_lines(x, plot[3] + 55, [f"{tick:.2g}"], 28, fill=COLORS["gray"]))
    x_ref = x_or(1.0)
    draw.line([(x_ref, plot[1]), (x_ref, plot[3])], fill="#808080", width=4)
    svg_parts.append(f'<line x1="{x_ref}" y1="{plot[1]}" x2="{x_ref}" y2="{plot[3]}" stroke="#808080" stroke-width="4" stroke-dasharray="12 12"/>')

    y_positions = [760, 1120, 1480, 1840]
    for y, (label, row, color) in zip(y_positions, rows):
        draw_multiline(draw, (label_x, y - 70, 1060, y + 70), label, font(34), fill="#111111", align="left")
        low, mid, high = x_or(row["ci_low"]), x_or(row["or"]), x_or(row["ci_high"])
        draw.line([(low, y), (high, y)], fill=color, width=8)
        draw.line([(low, y - 32), (low, y + 32)], fill=color, width=6)
        draw.line([(high, y - 32), (high, y + 32)], fill=color, width=6)
        draw.ellipse((mid - 22, y - 22, mid + 22, y + 22), fill=color, outline="white", width=4)
        estimate = f"{row['or']:.3f} ({row['ci_low']:.3f}-{row['ci_high']:.3f}); P={fmt_p(row['p_value'])}"
        draw.text((estimate_x, y - 22), estimate, font=font(33), fill="#111111")
        svg_parts.append(svg_wrapped_text(label_x, y - 38, label, 34, 28, fill="#111111", anchor="start"))
        svg_parts.append(f'<line x1="{low}" y1="{y}" x2="{high}" y2="{y}" stroke="{color}" stroke-width="8"/><line x1="{low}" y1="{y-32}" x2="{low}" y2="{y+32}" stroke="{color}" stroke-width="6"/><line x1="{high}" y1="{y-32}" x2="{high}" y2="{y+32}" stroke="{color}" stroke-width="6"/><circle cx="{mid}" cy="{y}" r="22" fill="{color}" stroke="white" stroke-width="4"/>')
        svg_parts.append(svg_text_lines(estimate_x, y + 11, [estimate], 33, anchor="start"))

    draw.text((plot[0], plot[3] + 85), "Odds ratio per eICU SD of hypotension burden", font=font(34), fill="#111111")
    svg_parts.append(svg_text_lines(plot[0], plot[3] + 130, ["Odds ratio per eICU SD of hypotension burden"], 34, anchor="start"))

    cal = src.v1_cal[src.v1_cal["dataset"].eq("MIMIC_IV_external_validation")].sort_values("risk_group")
    draw.text((3050, 310), "B. External validation calibration", font=font(44, True), fill=COLORS["navy"])
    svg_parts.append(svg_text_lines(3050, 350, ["B. External validation calibration"], 44, fill=COLORS["navy"], weight="bold", anchor="start"))
    cplot = (3180, 610, 4890, 2200)
    risk_min, risk_max = 0.0, 0.55

    def cx(value: float) -> int:
        return int(cplot[0] + (value - risk_min) / (risk_max - risk_min) * (cplot[2] - cplot[0]))

    def cy(value: float) -> int:
        return int(cplot[3] - (value - risk_min) / (risk_max - risk_min) * (cplot[3] - cplot[1]))

    draw.rectangle(cplot, outline="#D1D5DB", width=4)
    svg_parts.append(f'<rect x="{cplot[0]}" y="{cplot[1]}" width="{cplot[2]-cplot[0]}" height="{cplot[3]-cplot[1]}" fill="none" stroke="#D1D5DB" stroke-width="4"/>')
    for tick in np.linspace(0, 0.5, 6):
        x = cx(float(tick))
        y = cy(float(tick))
        draw.line([(x, cplot[1]), (x, cplot[3])], fill="#E5E7EB", width=3)
        draw.line([(cplot[0], y), (cplot[2], y)], fill="#E5E7EB", width=3)
        draw.text((x - 22, cplot[3] + 24), f"{tick:.1f}", font=font(28), fill=COLORS["gray"])
        draw.text((cplot[0] - 72, y - 16), f"{tick:.1f}", font=font(28), fill=COLORS["gray"])
        svg_parts.append(f'<line x1="{x}" y1="{cplot[1]}" x2="{x}" y2="{cplot[3]}" stroke="#E5E7EB" stroke-width="3"/><line x1="{cplot[0]}" y1="{y}" x2="{cplot[2]}" y2="{y}" stroke="#E5E7EB" stroke-width="3"/>')
        svg_parts.append(svg_text_lines(x, cplot[3] + 55, [f"{tick:.1f}"], 28, fill=COLORS["gray"]))
        svg_parts.append(svg_text_lines(cplot[0] - 45, y + 9, [f"{tick:.1f}"], 28, fill=COLORS["gray"], anchor="end"))
    draw.line([(cx(0), cy(0)), (cx(0.55), cy(0.55))], fill="#9CA3AF", width=5)
    svg_parts.append(f'<line x1="{cx(0)}" y1="{cy(0)}" x2="{cx(0.55)}" y2="{cy(0.55)}" stroke="#9CA3AF" stroke-width="5" stroke-dasharray="15 15"/>')
    points = [(cx(float(r["mean_predicted_risk"])), cy(float(r["observed_risk"]))) for _, r in cal.iterrows()]
    draw.line(points, fill=COLORS["blue"], width=7)
    svg_parts.append('<polyline points="' + " ".join(f"{x},{y}" for x, y in points) + f'" fill="none" stroke="{COLORS["blue"]}" stroke-width="7"/>')
    for (x, y), (_, row) in zip(points, cal.iterrows()):
        radius = int(np.clip(row["n"] / 2.2, 38, 82))
        draw.ellipse((x - radius, y - radius, x + radius, y + radius), fill=COLORS["blue"], outline="white", width=6)
        draw.text((x - 22, y - radius - 44), f"Q{int(row['risk_group'])}", font=font(28), fill=COLORS["gray"])
        svg_parts.append(f'<circle cx="{x}" cy="{y}" r="{radius}" fill="{COLORS["blue"]}" stroke="white" stroke-width="6"/>')
        svg_parts.append(svg_text_lines(x, y - radius - 18, [f"Q{int(row['risk_group'])}"], 28, fill=COLORS["gray"]))

    mimic_perf = perf(src.v1_perf)
    eicu_perf = perf(src.v1_perf, dataset="eICU_derivation")
    metric_text = (
        f"eICU C-index: {eicu_perf['auc_c_index']:.3f}\n"
        f"MIMIC-IV C-index: {mimic_perf['auc_c_index']:.3f}\n"
        f"MIMIC-IV Brier: {mimic_perf['brier_score']:.3f}\n"
        f"Calibration slope: {mimic_perf['calibration_slope_logistic']:.3f}\n"
        f"Calibration intercept: {mimic_perf['calibration_intercept_logistic']:.3f}"
    )
    metric_box = (3130, 2250, 4920, 2655)
    draw.rounded_rectangle(metric_box, radius=30, fill=COLORS["gray_light"], outline="#D1D5DB", width=4)
    draw_multiline(draw, (metric_box[0] + 40, metric_box[1] + 35, metric_box[2] - 40, metric_box[3] - 35), metric_text, font(33), fill="#111111", align="left")
    draw.text((3700, cplot[3] + 95), "Mean predicted risk", font=font(34), fill="#111111")
    draw.text((2915, 1380), "Observed risk", font=font(34), fill="#111111")
    svg_parts.append(f'<rect x="{metric_box[0]}" y="{metric_box[1]}" width="{metric_box[2]-metric_box[0]}" height="{metric_box[3]-metric_box[1]}" rx="30" fill="{COLORS["gray_light"]}" stroke="#D1D5DB" stroke-width="4"/>')
    svg_parts.append(svg_wrapped_text(metric_box[0] + 45, metric_box[1] + 78, metric_text, 33, 55, anchor="start"))
    svg_parts.append(svg_text_lines(4050, cplot[3] + 140, ["Mean predicted risk"], 34))
    svg_parts.append(svg_text_lines(2915, 1415, ["Observed risk"], 34, anchor="start"))

    foot = "Primary model adjusted for hypoxemia burden, age, sex, and GCS. Sensitivity rows use the same covariate set unless otherwise specified."
    draw_multiline(draw, (360, 2880, W - 360, 2960), foot, font(30), fill=COLORS["gray"])
    svg_parts.append(svg_wrapped_text(W / 2, 2930, foot, 30, 150, fill=COLORS["gray"]))
    svg_parts.append("</svg>")
    return save_svg_png_tiff(img, "\n".join(svg_parts), "Figure_2_Main_Effect_External_Validation", dpi=600)


def build_supplementary_rcs_figure(src: ArtifactSources) -> dict[str, str]:
    from supplementary_diagnostics_figures_v1 import build_rcs_figure

    draft_paths = build_rcs_figure()
    svg_path = FIGURE_DIR / "Supplementary_Figure_S1_RCS_Exploratory.svg"
    png_path = FIGURE_DIR / "Supplementary_Figure_S1_RCS_Exploratory_600dpi.png"
    tiff_path = FIGURE_DIR / "Supplementary_Figure_S1_RCS_Exploratory_600dpi.tiff"

    shutil.copyfile(draft_paths["svg"], svg_path)
    shutil.copyfile(draft_paths["tiff"], tiff_path)
    with Image.open(draft_paths["png"]) as preview:
        pixel_size = f"{preview.width}x{preview.height}"
        preview.save(png_path, dpi=(600, 600))

    return {
        "svg": str(svg_path),
        "png": str(png_path),
        "tiff": str(tiff_path),
        "pixel_size": pixel_size,
        "dpi": "600",
    }

    W, H = 4200, 3000
    img = Image.new("RGB", (W, H), "white")
    draw = ImageDraw.Draw(img)
    original = src.rcs_curve[src.rcs_curve["model"].eq("hypotension_rcs_3k_primary")].copy()
    winsor = src.winsor_curve.copy()
    x_cap = float(src.winsor_coef["cap_value_eicu_p99"].dropna().iloc[0])
    original = original[original["hypotension_twa"].le(x_cap)]
    winsor = winsor[winsor["hypotension_twa"].le(x_cap)]
    ymax = float(max(original["odds_ratio_vs_zero"].max(), winsor["odds_ratio_vs_zero"].max()) * 1.08)
    ymin = 0.85
    plot = (520, 470, 3850, 2290)

    def x_map(v: float) -> int:
        return int(plot[0] + v / x_cap * (plot[2] - plot[0]))

    def y_map(v: float) -> int:
        return int(plot[3] - (v - ymin) / (ymax - ymin) * (plot[3] - plot[1]))

    draw.text((360, 170), "Exploratory shape of the hypotension burden association", font=font(56, True), fill=COLORS["navy"])
    draw.rectangle(plot, outline="#D1D5DB", width=4)
    svg_parts = [
        f'<svg xmlns="http://www.w3.org/2000/svg" width="{W}" height="{H}" viewBox="0 0 {W} {H}">',
        '<rect width="100%" height="100%" fill="white"/>',
        svg_text_lines(360, 225, ["Exploratory shape of the hypotension burden association"], 56, fill=COLORS["navy"], weight="bold", anchor="start"),
        f'<rect x="{plot[0]}" y="{plot[1]}" width="{plot[2]-plot[0]}" height="{plot[3]-plot[1]}" fill="none" stroke="#D1D5DB" stroke-width="4"/>',
    ]
    for tick in np.linspace(0, x_cap, 6):
        x = x_map(float(tick))
        draw.line([(x, plot[1]), (x, plot[3])], fill="#E5E7EB", width=3)
        draw.text((x - 28, plot[3] + 28), f"{tick:.1f}", font=font(30), fill=COLORS["gray"])
        svg_parts.append(f'<line x1="{x}" y1="{plot[1]}" x2="{x}" y2="{plot[3]}" stroke="#E5E7EB" stroke-width="3"/>')
        svg_parts.append(svg_text_lines(x, plot[3] + 62, [f"{tick:.1f}"], 30, fill=COLORS["gray"]))
    for tick in np.linspace(1, math.ceil(ymax), max(2, math.ceil(ymax))):
        y = y_map(float(tick))
        draw.line([(plot[0], y), (plot[2], y)], fill="#E5E7EB", width=3)
        draw.text((plot[0] - 85, y - 18), f"{tick:.0f}", font=font(30), fill=COLORS["gray"])
        svg_parts.append(f'<line x1="{plot[0]}" y1="{y}" x2="{plot[2]}" y2="{y}" stroke="#E5E7EB" stroke-width="3"/>')
        svg_parts.append(svg_text_lines(plot[0] - 50, y + 10, [f"{tick:.0f}"], 30, fill=COLORS["gray"], anchor="end"))
    y1 = y_map(1)
    draw.line([(plot[0], y1), (plot[2], y1)], fill="#9CA3AF", width=4)
    draw.line([(x_map(x_cap), plot[1]), (x_map(x_cap), plot[3])], fill="#D1D5DB", width=4)
    svg_parts.append(f'<line x1="{plot[0]}" y1="{y1}" x2="{plot[2]}" y2="{y1}" stroke="#9CA3AF" stroke-width="4" stroke-dasharray="12 12"/>')

    def curve_points(df: pd.DataFrame) -> list[tuple[int, int]]:
        return [(x_map(float(r["hypotension_twa"])), y_map(float(r["odds_ratio_vs_zero"]))) for _, r in df.iterrows()]

    orig_pts = curve_points(original)
    win_pts = curve_points(winsor)
    draw.line(orig_pts, fill=COLORS["purple"], width=10, joint="curve")
    draw.line(win_pts, fill=COLORS["gold"], width=10, joint="curve")
    svg_parts.append('<polyline points="' + " ".join(f"{x},{y}" for x, y in orig_pts) + f'" fill="none" stroke="{COLORS["purple"]}" stroke-width="10"/>')
    svg_parts.append('<polyline points="' + " ".join(f"{x},{y}" for x, y in win_pts) + f'" fill="none" stroke="{COLORS["gold"]}" stroke-width="10" stroke-dasharray="24 18"/>')

    draw.line([(2670, 430), (2860, 430)], fill=COLORS["purple"], width=10)
    draw.text((2890, 408), "Original 3-knot RCS", font=font(32), fill="#111111")
    draw.line([(2670, 500), (2860, 500)], fill=COLORS["gold"], width=10)
    draw.text((2890, 478), "99% winsorized 3-knot RCS", font=font(32), fill="#111111")
    svg_parts.append(f'<line x1="2670" y1="430" x2="2860" y2="430" stroke="{COLORS["purple"]}" stroke-width="10"/>{svg_text_lines(2890, 438, ["Original 3-knot RCS"], 32, anchor="start")}')
    svg_parts.append(f'<line x1="2670" y1="500" x2="2860" y2="500" stroke="{COLORS["gold"]}" stroke-width="10"/>{svg_text_lines(2890, 508, ["99% winsorized 3-knot RCS"], 32, anchor="start")}')

    annotation = "Original RCS vs linear P=0.019\nWinsorized RCS vs linear P=0.132\nValues are not clinical cutoffs."
    ann_box = (640, 650, 1840, 960)
    draw.rounded_rectangle(ann_box, radius=28, fill=COLORS["purple_light"], outline=COLORS["purple"], width=4)
    draw_multiline(draw, (ann_box[0] + 35, ann_box[1] + 30, ann_box[2] - 35, ann_box[3] - 30), annotation, font(32), align="left")
    svg_parts.append(f'<rect x="{ann_box[0]}" y="{ann_box[1]}" width="{ann_box[2]-ann_box[0]}" height="{ann_box[3]-ann_box[1]}" rx="28" fill="{COLORS["purple_light"]}" stroke="{COLORS["purple"]}" stroke-width="4"/>')
    svg_parts.append(svg_wrapped_text(ann_box[0] + 35, ann_box[1] + 75, annotation, 32, 40, anchor="start"))

    x_label = "Hypotension burden (MAP <65 time-weighted deficit)"
    y_label = "Adjusted odds ratio vs zero burden"
    draw_multiline(draw, (780, 2400, 3600, 2470), x_label, font(34), fill="#111111")
    draw.text((75, 1350), y_label, font=font(34), fill="#111111")
    foot = "Adjusted at reference covariate values from the eICU model. The winsorized check attenuated formal nonlinearity, so the curve is treated as exploratory."
    draw_multiline(draw, (320, 2850, W - 320, 2940), foot, font(30), fill=COLORS["gray"])
    svg_parts.append(svg_text_lines((plot[0] + plot[2]) / 2, 2450, [x_label], 34))
    svg_parts.append(svg_text_lines(80, 1390, [y_label], 34, anchor="start"))
    svg_parts.append(svg_wrapped_text(W / 2, 2900, foot, 30, 145, fill=COLORS["gray"]))
    svg_parts.append("</svg>")
    return save_svg_png_tiff(img, "\n".join(svg_parts), "Supplementary_Figure_S1_RCS_Exploratory", dpi=600)


def write_tables(src: ArtifactSources) -> tuple[pd.DataFrame, pd.DataFrame, dict[str, pd.DataFrame], list[Path]]:
    table1 = build_table1(src)
    table2 = build_table2(src)
    supp = build_supplementary_tables(src)

    table1.to_csv(TABLE_DIR / "Table_1_Baseline_Characteristics.csv", index=False, encoding="utf-8-sig")
    table2.to_csv(TABLE_DIR / "Table_2_Main_Model_And_Sensitivity_Analyses.csv", index=False, encoding="utf-8-sig")
    write_table1_markdown(
        TABLE_DIR / "Table_1_Baseline_Characteristics.md",
        "Table 1. Mixed baseline characteristics of the landmark severe TBI cohorts",
        table1,
        "Continuous variables are median [IQR]; categorical variables are n (%). P values compare ICU survivors and ICU deaths within each database. Hospital death is shown last as a descriptive outcome descriptor; no P value is reported for that row because ICU death defines the stratification. First-24h mechanical ventilation and vasopressor use are descriptive treatment-window variables and were not included in the primary model because they overlap the exposure window and may represent treatment response or mediation. MIMIC-IV herniation counts reflect source coding and should not be interpreted as true absence or as a cross-database subtype contrast. Sex percentages use non-missing denominators; sex was missing for 2 eICU survivors.",
    )
    write_markdown_table(
        TABLE_DIR / "Table_2_Main_Model_And_Sensitivity_Analyses.md",
        "Table 2. Primary model and prespecified sensitivity analyses",
        table2,
        "Effect estimates are odds ratios per eICU SD unless otherwise stated.",
    )

    for name, df in supp.items():
        df.to_csv(TABLE_DIR / f"{name}.csv", index=False, encoding="utf-8-sig")
        write_markdown_table(TABLE_DIR / f"{name}.md", name.replace("_", " "), df)

    docx_paths = write_docx_tables(table1, table2, supp)
    return table1, table2, supp, docx_paths


def write_decision_note(src: ArtifactSources) -> Path:
    original_p = src.rcs_lr[src.rcs_lr["comparison"].eq("hypotension_rcs_3k_primary_vs_linear_hypotension")]["p_value"].iloc[0]
    winsor_p = src.winsor_lr[src.winsor_lr["comparison"].eq("winsor99_3k_rcs_vs_linear")]["p_value"].iloc[0]
    note = f"""# Final Artifact Decision Note

Run ID: `{RUN_ID}`

## Primary Decision

The final manuscript figures and tables use the linear hypotension-burden main effect as the core result.

Rationale:

- The adjusted linear model is robust: primary ICU death OR per eICU SD = 1.350, 95% CI 1.120-1.629, P = 0.0017.
- Sensitivity analyses remain directionally consistent: hospital death, MAP <60, and 99% winsorization all support the main association.
- The original 3-knot RCS suggested nonlinearity (RCS vs linear P = {original_p:.4f}), but this attenuated after 99th-percentile winsorization (P = {winsor_p:.4f}).

Therefore, RCS is retained as main-text Figure 3 and Supplementary Table S6, with wording limited to exploratory shape instability rather than clinical cutoffs.

## Final Artifact Inventory

- Figure 1: Study flow.
- Figure 2: Main hypotension effect and MIMIC-IV external validation.
- Figure 3: Exploratory RCS shape robustness with winsorized tail-sensitivity check.
- Table 1: Baseline characteristics.
- Table 2: Main model and sensitivity analyses.
- Supplementary Tables S1-S9: cohort flow, exclusions, QC/missingness, zero inflation, interactions, RCS tests, validation performance, collinearity, and correlations.
"""
    path = OUTPUT_ROOT / "FINAL_DECISION_NOTE.md"
    path.write_text(note, encoding="utf-8")
    return path


def write_manifest(artifacts: dict[str, object]) -> Path:
    manifest = {
        "run_id": RUN_ID,
        "project_root": str(PROJECT_ROOT),
        "output_root": str(OUTPUT_ROOT),
        "artifacts": artifacts,
        "source_outputs": {
            "final_covariates": "03_outputs/tables/covariates/20260620_covariates_v2_baseline",
            "primary_model": "03_outputs/tables/models/20260620_adjusted_model_v1",
            "interaction_model": "03_outputs/tables/models/20260620_adjusted_model_v2_interaction",
            "rcs_model": "03_outputs/tables/models/20260620_adjusted_model_v3b_hypotension_rcs_limited",
            "sensitivity": "03_outputs/tables/sensitivity/20260621_analysis_closure_v1",
            "qc": "03_outputs/tables/qc/20260620_final_dataset_qc_v1",
        },
    }
    path = OUTPUT_ROOT / "manifest.json"
    path.write_text(json.dumps(manifest, indent=2, ensure_ascii=False), encoding="utf-8")
    return path


def write_qa_summary(src: ArtifactSources, figure_artifacts: dict[str, dict[str, str]], docx_paths: list[Path]) -> Path:
    fail_rules = src.rule_checks[src.rule_checks["status"].eq("FAIL")]
    warn_rules = src.rule_checks[src.rule_checks["status"].eq("WARN")]
    lines = [
        "# Final Artifact QA Summary",
        "",
        f"Run ID: `{RUN_ID}`",
        "",
        "## Numeric Checks",
        "",
        f"- Final eICU landmark N = {len(src.eicu)}; ICU deaths = {int(src.eicu['death_icu'].sum())}; hospital deaths = {int(src.eicu['death_hospital'].sum())}.",
        f"- Final MIMIC-IV landmark N = {len(src.mimic)}; ICU deaths = {int(src.mimic['death_icu'].sum())}; hospital deaths = {int(src.mimic['death_hospital'].sum())}.",
        f"- QC FAIL rules = {len(fail_rules)}.",
        f"- QC WARN rules = {len(warn_rules)}; expected eICU sex_male missingness warning is retained.",
        "",
        "## Visual/Export Checks",
        "",
    ]
    for name, artifact in figure_artifacts.items():
        lines.append(f"- {name}: PNG {artifact['pixel_size']} at {artifact['dpi']} dpi; SVG/PNG/TIFF exported.")
    lines += [
        "",
        "## DOCX Tables",
        "",
    ]
    for path in docx_paths:
        lines.append(f"- {path.name}")
    lines += [
        "",
        "DOCX visual rendering is performed separately after generation when LibreOffice/Word rendering is available.",
    ]
    path = QA_DIR / "final_artifact_qa_summary.md"
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return path


def main() -> None:
    ensure_dirs()
    src = load_sources()
    table1, table2, supp, docx_paths = write_tables(src)

    figures = {
        "Figure 1": build_figure1(src),
        "Figure 2": build_figure2(src),
        "Supplementary Figure S1": build_supplementary_rcs_figure(src),
    }

    decision_note = write_decision_note(src)
    qa_summary = write_qa_summary(src, figures, docx_paths)
    manifest = write_manifest({
        "tables": [str(p) for p in TABLE_DIR.glob("*")],
        "figures": figures,
        "docx": [str(p) for p in docx_paths],
        "decision_note": str(decision_note),
        "qa_summary": str(qa_summary),
    })

    print(json.dumps({
        "run_id": RUN_ID,
        "output_root": str(OUTPUT_ROOT),
        "table1_rows": len(table1),
        "table2_rows": len(table2),
        "supplementary_tables": len(supp),
        "docx_files": [str(p) for p in docx_paths],
        "figures": figures,
        "manifest": str(manifest),
    }, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt

    main()
