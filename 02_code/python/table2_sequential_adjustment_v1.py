from __future__ import annotations

import csv
from pathlib import Path

import numpy as np

import adjusted_model_v1 as base


PROJECT_ROOT = Path(__file__).resolve().parents[2]
FINAL_TABLE_DIR = (
    PROJECT_ROOT
    / "03_outputs"
    / "final_artifacts"
    / "20260621_final_manuscript_artifacts_v1"
    / "tables"
)

CSV_PATH = FINAL_TABLE_DIR / "Table_2_Sequential_Adjustment_Models.csv"
MD_PATH = FINAL_TABLE_DIR / "Table_2_Sequential_Adjustment_Models.md"
DOCX_PATH = FINAL_TABLE_DIR / "Table_2_Sequential_Adjustment_Models.docx"


TITLE = "Table 2. Sequential adjustment models for the association between hypotension burden and ICU death"
NOTE = (
    "Note: Effect estimates are odds ratios for ICU death after the 24-h landmark per eICU SD increase in "
    "MAP <65 mmHg hypotension burden. All rows use the same eICU complete-case analysis set. C-index and "
    "Brier score are apparent eICU performance metrics and are not optimism-corrected; external validation "
    "of the fully adjusted model in MIMIC-IV is shown in Figure 2 and Supplementary Table S7. The Brier "
    "score is strongly influenced by the baseline event rate and is less sensitive than the C-index to "
    "improvements in rank discrimination."
)


MODEL_SPECS = [
    (
        "Unadjusted model",
        "Hypotension burden only",
        ["intercept", "hypotension_twa_per_sd"],
    ),
    (
        "Clinical adjustment",
        "Hypotension burden + age + sex + GCS",
        ["intercept", "hypotension_twa_per_sd", "age_per_10y", "sex_male", "gcs_per_point"],
    ),
    (
        "Fully adjusted primary model",
        "Hypotension burden + hypoxemia burden + age + sex + GCS",
        [
            "intercept",
            "hypotension_twa_per_sd",
            "hypoxemia_twa_per_sd",
            "age_per_10y",
            "sex_male",
            "gcs_per_point",
        ],
    ),
]


def fmt_p(value: float) -> str:
    if value < 0.001:
        return "<0.001"
    return f"{value:.3f}"


def build_rows() -> list[dict[str, str]]:
    df = base.read_analysis(base.EICU_PATH)
    model_df = df[base.CORE_COLUMNS].dropna().copy()
    params = base.build_transform_params(model_df)
    feature_df = base.transform_features(model_df, params)
    y = model_df[base.OUTCOME].to_numpy(dtype=float)

    rows: list[dict[str, str]] = []
    for model_name, adjustment, columns in MODEL_SPECS:
        fit = base.fit_logistic(feature_df[columns].to_numpy(dtype=float), y)
        coef = base.coefficient_table(columns, fit)
        hypo = coef.loc[coef["variable"] == "hypotension_twa_per_sd"].iloc[0]
        pred = fit["predicted"]
        rows.append(
            {
                "Model": model_name,
                "Adjustment set": adjustment,
                "eICU analysis set (N/events)": f"{len(y)}/{int(y.sum())}",
                "OR (95% CI)": f"{hypo['or']:.3f} ({hypo['ci_low']:.3f}-{hypo['ci_high']:.3f})",
                "P value": fmt_p(float(hypo["p_value"])),
                "eICU apparent C-index": f"{base.auc_score(y, pred):.3f}",
                "Brier score": f"{float(np.mean((y - pred) ** 2)):.3f}",
            }
        )
    return rows


def write_csv(rows: list[dict[str, str]]) -> None:
    FINAL_TABLE_DIR.mkdir(parents=True, exist_ok=True)
    with CSV_PATH.open("w", encoding="utf-8", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=list(rows[0].keys()))
        writer.writeheader()
        writer.writerows(rows)


def write_md(rows: list[dict[str, str]]) -> None:
    headers = list(rows[0].keys())
    lines = [f"# {TITLE}", ""]
    lines.append("| " + " | ".join(headers) + " |")
    lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
    for row in rows:
        lines.append("| " + " | ".join(row[h] for h in headers) + " |")
    lines.extend(["", NOTE, ""])
    MD_PATH.write_text("\n".join(lines), encoding="utf-8")


def set_cell_text(cell, text: str, bold: bool = False, size: float = 9.5) -> None:
    from docx.shared import Pt

    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = 0
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)


def set_cell_border(cell, edge: str, val: str = "nil", size: int = 0) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    tag = qn(f"w:{edge}")
    element = tc_borders.find(tag)
    if element is None:
        element = OxmlElement(f"w:{edge}")
        tc_borders.append(element)

    element.set(qn("w:val"), val)
    if val != "nil":
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def apply_three_line_table_borders(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            for edge in ("top", "bottom", "left", "right"):
                set_cell_border(cell, edge, "nil")

    for cell in table.rows[0].cells:
        set_cell_border(cell, "top", "single", size=12)
        set_cell_border(cell, "bottom", "single", size=8)

    for cell in table.rows[-1].cells:
        set_cell_border(cell, "bottom", "single", size=12)


def write_docx(rows: list[dict[str, str]]) -> None:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.shared import Cm, Pt

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(10)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(6)
    title_run = title.add_run(TITLE)
    title_run.bold = True
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(11)

    headers = list(rows[0].keys())
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    widths_cm = [3.1, 5.2, 3.0, 3.2, 1.8, 2.5, 2.0]
    for col_idx, width in enumerate(widths_cm):
        for cell in table.columns[col_idx].cells:
            cell.width = Cm(width)

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_text(cell, header, bold=True, size=9)

    for row in rows:
        cells = table.add_row().cells
        for idx, header in enumerate(headers):
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_text(cells[idx], row[header], bold=False, size=9)

    apply_three_line_table_borders(table)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(6)
    note.paragraph_format.space_after = Pt(0)
    note_run = note.add_run(NOTE)
    note_run.font.name = "Times New Roman"
    note_run.font.size = Pt(8.5)

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    from docx.shared import Pt

    rows = build_rows()
    write_csv(rows)
    write_md(rows)
    write_docx(rows)
    print(CSV_PATH)
    print(MD_PATH)
    print(DOCX_PATH)
